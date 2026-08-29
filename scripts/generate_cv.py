#!/usr/bin/env python3
"""
Generate CV for AI Platform Applications using python-docx
Template A: Left Sidebar + Right Body
"""
import sys, os

# Check for python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx -q")
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

OUTPUT_DIR = '/home/z/my-project/download'
OUTPUT_FILE = os.path.join(OUTPUT_DIR, 'Kennedy_Hachimeganum_CV.docx')

# ─── Color Palette (Template A) ───
class S:
    bg = RGBColor(0x3B, 0x4F, 0x5C)        # sidebar
    text = RGBColor(0xD8, 0xE2, 0xE8)       # sidebar text
    label = RGBColor(0x8B, 0xA0, 0xAD)      # sidebar secondary
    accent = RGBColor(0x2F, 0x97, 0xB8)     # accent blue
    title = RGBColor(0x1A, 0x2D, 0x38)      # body heading
    body = RGBColor(0x2C, 0x3E, 0x4A)       # body text
    sec = RGBColor(0x6B, 0x85, 0x92)        # secondary info


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("style", "single")}" '
            f'w:sz="{val.get("size", "4")}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_paragraph_to_cell(cell, text, font_name='Times New Roman', font_size=11,
                          bold=False, italic=False, color=None, alignment=None,
                          space_before=0, space_after=0, line_spacing=1.15):
    """Add a formatted paragraph to a table cell."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_sidebar_section(cell, label, value, label_color=None):
    """Add a labeled section in the sidebar."""
    lc = label_color or S.label
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8.5)
    run.font.color.rgb = lc
    run.font.bold = True

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run(value)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(9.5)
    run2.font.color.rgb = S.text


def add_skill_rating(cell, skill_name, level, detail=''):
    """Add a skill with dot rating to sidebar."""
    filled = level
    empty = 5 - level
    dots = '\u25CF ' * filled + '\u25CB ' * empty
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f'{skill_name}  ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    run.font.color.rgb = S.text
    run.font.bold = True
    run2 = p.add_run(dots)
    run2.font.name = 'Segoe UI Symbol'
    run2.font.size = Pt(7)
    run2.font.color.rgb = S.accent
    if detail:
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(2)
        run3 = p3.add_run(detail)
        run3.font.name = 'Times New Roman'
        run3.font.size = Pt(8)
        run3.font.color.rgb = S.label


def build_cv():
    doc = Document()

    # ─── Page setup: zero margins (table controls layout) ───
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    # ─── Main table: sidebar + body ───
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths: 30% sidebar + 70% body
    sidebar_width = Cm(6.0)
    body_width = Cm(15.0)
    for cell in table.columns[0].cells:
        cell.width = sidebar_width
    for cell in table.columns[1].cells:
        cell.width = body_width

    sidebar_cell = table.cell(0, 0)
    body_cell = table.cell(0, 1)

    # ─── SIDEBAR ───
    set_cell_shading(sidebar_cell, '3B4F5C')
    sidebar_cell.vertical_alignment = 1  # CENTER

    # Remove default paragraph
    for p in sidebar_cell.paragraphs:
        p.clear()

    # Name
    add_paragraph_to_cell(sidebar_cell, 'WIKE-YOUNG\nKENNEDY\nHACHIMEGANUM',
                         font_name='Times New Roman', font_size=22,
                         bold=True, color=S.text, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=20, space_after=4)

    # Divider line (using paragraph border)
    div_p = sidebar_cell.add_paragraph()
    div_p.paragraph_format.space_before = Pt(4)
    div_p.paragraph_format.space_after = Pt(8)
    div_pPr = div_p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="2F97B8"/>'
        f'</w:pBdr>'
    )
    div_pPr.append(pBdr)

    # Target Position
    add_paragraph_to_cell(sidebar_cell, 'AI Data Annotator\n& Remote Contributor',
                         font_name='Times New Roman', font_size=12,
                         bold=False, italic=True, color=S.accent,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=0, space_after=12)

    # Basic Info
    add_sidebar_section(sidebar_cell, 'LOCATION', 'Port Harcourt, Rivers State, Nigeria')

    # Contact
    add_sidebar_section(sidebar_cell, 'CONTACT', '')
    add_paragraph_to_cell(sidebar_cell, 'wikeyoung41@gmail.com',
                         font_size=9, color=S.text, space_before=1, space_after=1)
    add_paragraph_to_cell(sidebar_cell, '+234 816 012 4516',
                         font_size=9, color=S.text, space_before=0, space_after=1)

    # Languages
    add_sidebar_section(sidebar_cell, 'LANGUAGES', '')
    add_skill_rating(sidebar_cell, 'English', 4, 'Fluent (reading, writing, speaking)')
    add_skill_rating(sidebar_cell, 'Ikwerre', 5, 'Native speaker')
    add_skill_rating(sidebar_cell, 'Igbo', 2, 'Partial comprehension')

    # Skills
    add_sidebar_section(sidebar_cell, 'SKILLS', '')
    add_skill_rating(sidebar_cell, 'Data Annotation', 3, 'Image labeling, text classification')
    add_skill_rating(sidebar_cell, 'Writing & Comprehension', 4, 'Clear written communication')
    add_skill_rating(sidebar_cell, 'Internet Research', 4, 'Search evaluation, fact-checking')
    add_skill_rating(sidebar_cell, 'Attention to Detail', 4, 'Guideline adherence, quality focus')
    add_skill_rating(sidebar_cell, 'Typing Speed', 2, 'Currently improving (30+ WPM)')
    add_skill_rating(sidebar_cell, 'Time Management', 3, 'Consistent daily scheduling')

    # Education
    add_sidebar_section(sidebar_cell, 'EDUCATION', '')
    add_paragraph_to_cell(sidebar_cell, 'O\'Level Certificate',
                         font_size=9.5, bold=True, color=S.text,
                         space_before=2, space_after=0)
    add_paragraph_to_cell(sidebar_cell, 'Secondary Education completed.',
                         font_size=8.5, color=S.label,
                         space_before=0, space_after=4)

    # Interests
    add_sidebar_section(sidebar_cell, 'INTERESTS', '')
    add_paragraph_to_cell(sidebar_cell,
        'AI technology, language preservation, content evaluation, online learning, self-development',
        font_size=8.5, color=S.text, space_before=1, space_after=4)

    # ─── BODY ───
    body_cell.vertical_alignment = 1  # TOP

    # Remove default paragraph
    for p in body_cell.paragraphs:
        p.clear()

    body_margin_left = Cm(0.8)
    body_margin_right = Cm(0.6)

    # ─── Helper: Section heading with blue bar ───
    def add_section_heading(text, subtitle=''):
        # Use a single-cell table for the colored bar
        heading_table = body_cell.add_table(rows=1, cols=1)
        heading_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hc = heading_table.cell(0, 0)
        set_cell_shading(hc, '2F97B8')
        hc.width = body_width - body_margin_left - body_margin_right
        # Remove borders
        for edge in ['top', 'left', 'bottom', 'right']:
            set_cell_border(hc, **{edge: {'style': 'none', 'size': '0'}})
        hp = hc.paragraphs[0]
        hp.paragraph_format.space_before = Pt(10)
        hp.paragraph_format.space_after = Pt(6)
        # Main heading text
        run = hp.add_run(f'  {text}  ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if subtitle:
            run2 = hp.add_run(f'  {subtitle}')
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(9)
            run2.font.color.rgb = RGBColor(0xC8, 0xE8, 0xF0)
            run2.font.italic = True

    def add_bullet(text, indent=0.5):
        p = body_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(indent)
        run = p.add_run(f'\u25B8  {text}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9.5)
        run.font.color.rgb = S.body

    def add_body_text(text, bold=False, italic=False, space_before=6, space_after=4):
        p = body_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.color.rgb = S.body
        run.font.bold = bold
        run.font.italic = italic

    # ─── PROFILE SUMMARY ───
    add_section_heading('Profile Summary')
    add_body_text(
        'Detail-oriented and highly motivated individual seeking AI data annotation and remote '
        'contributor roles. Native Ikwerre speaker with strong English proficiency, providing a '
        'unique advantage for multilingual AI training and localization projects. Reliable internet '
        'access in Port Harcourt with commitment to consistent daily work schedules. Focused on '
        'building a sustainable remote income through quality-driven AI training contributions.',
        space_before=8, space_after=10
    )

    # ─── CORE COMPETENCIES ───
    add_section_heading('Core Competencies')
    comp_table = body_cell.add_table(rows=2, cols=4)
    comp_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    competencies = [
        'Text Classification & Evaluation',
        'Search Relevance Assessment',
        'Multilingual Data Annotation',
        'Image Labeling & Verification',
        'AI Output Quality Review',
        'Guideline Adherence',
        'Attention to Detail',
        'Consistent Daily Output'
    ]
    for i, comp in enumerate(competencies):
        row = i // 4
        col = i % 4
        c = comp_table.cell(row, col)
        cp = c.paragraphs[0]
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after = Pt(2)
        run = cp.add_run(f'\u2022  {comp}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8.5)
        run.font.color.rgb = S.body
        set_cell_border(c, **{'top': {'style': 'none'}, 'left': {'style': 'none'},
                               'bottom': {'style': 'none'}, 'right': {'style': 'none'}})

    # ─── RELEVANT SKILLS & CAPABILITIES ───
    add_section_heading('Relevant Skills & Capabilities')
    add_bullet('Strong reading comprehension and written English communication skills, essential for evaluating AI-generated text quality and providing corrective feedback.')
    add_bullet('Native fluency in Ikwerre language, enabling contribution to Nigerian-language AI training datasets that are in high demand for localization projects.')
    add_bullet('Partial comprehension of Igbo, expanding the range of linguistic tasks available across Welocalize, TELUS Digital, and CrowdGen localization projects.')
    add_bullet('Experience with internet-based research and information verification, directly applicable to search evaluation and fact-checking tasks on TELUS Digital and CrowdGen.')
    add_bullet('Demonstrated ability to follow detailed instructions and guidelines consistently, a critical requirement for maintaining quality scores above 90% on all AI training platforms.')
    add_bullet('Self-motivated with disciplined time management, capable of committing to 3-5 hours of daily platform work as required for consistent earnings.')

    # ─── PROJECTS & SELF-DEVELOPMENT ───
    add_section_heading('Projects & Self-Development')
    add_body_text('AI Platform Research & Preparation  |  2026', bold=True, italic=True, space_before=8)
    add_bullet('Completed comprehensive research on 9 AI data annotation platforms, analyzing legitimacy, Nigeria availability, payment methods, and earning potential.')
    add_bullet('Developed structured 90-day action plan for building sustainable remote income through AI training contributions across multiple platforms.')
    add_bullet('Currently improving typing speed through daily practice sessions, targeting 40+ WPM within 60 days for increased task throughput and earnings.')

    # ─── EDUCATION DETAIL ───
    add_section_heading('Education')
    add_body_text('O\'Level Certificate  |  Secondary Education', bold=True, italic=True, space_before=8)
    add_bullet('Completed secondary education with focus on English Language and general studies.')
    add_bullet('Strong foundational literacy and numeracy skills, continuously improved through self-study and online learning platforms.')

    # ─── AVAILABILITY ───
    add_section_heading('Availability')
    add_bullet('Available for immediate start on all Tier 1 platforms: CrowdGen (Appen), TELUS Digital, and Prolific.')
    add_bullet('Can commit 3-5 hours daily for consistent task completion and quality maintenance.')
    add_bullet('Flexible schedule with availability during peak task release times (early morning and evening hours).')
    add_bullet('Reliable internet connection and dedicated workspace in Port Harcourt, Rivers State.')

    # ─── REFEREES ───
    add_section_heading('References')
    add_body_text(
        'Available upon request. Professional references will be established through platform '
        'quality scores and contributor ratings as work progresses.',
        italic=True, space_before=6
    )

    # ─── Set row height to prevent blank page overflow ───
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="16038" w:hRule="exact"/>')
    trPr.append(trHeight)

    # ─── Save ───
    doc.save(OUTPUT_FILE)
    print(f'CV saved to: {OUTPUT_FILE}')
    return OUTPUT_FILE


if __name__ == '__main__':
    build_cv()
